# requires python-docx and python-dateutil

# simple modification of a word .docx template containing these form items:

# [CLIENT_NAME]
# [CLIENT_ADDRESS]
# [CLIENT_VAT]

# [INVOICE_NUMBER]
# [DATE]
# [DATE_DUE]

# [ITEMS_TABLE]

# Subtotal excluding tax: [SUBTOTAL]
# VAT: [VAT]
# Total: [TOTAL]

# [NOTES]


from docx import Document
from docx.shared import Pt, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from datetime import date
from dateutil.relativedelta import relativedelta

import invoicing  # local
from invoicing import RED, RESET, format_currency as fmtc
from decimal import Decimal

def fill_word_invoice(template_path, output_path, data, params = None):
    doc = Document(template_path)

    def replace_placeholder(paragraphs, replacements):
        for para in paragraphs:
            for key, val in replacements.items():
                if key in para.text:
                    inline = para.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            inline[i].text = inline[i].text.replace(key, val)

    default_params = {
        'currency_as_prefix': False,  # if true, prepend rather than append currency name/symbol
        'table_header_font': 'Libre Baskerville',
        'table_header_fontsize': 11,
        'table_header_fillcolor': '#101356',
        'table_vat_header': 'VAT',  # 'Sales Tax'
        'table_item_font': 'Calibri',
        'table_item_fontsize': 11,
        'table_bottomline_color': 'auto',
    }
    params = default_params | (params if params else {})  # overrides any or all default parameters
    cur1st = params['currency_as_prefix']

    # Replace table placeholder with real items table
    pretax_total = Decimal(0.0)
    total = Decimal(0.0)
    for i, para in enumerate(doc.paragraphs):
        if '[ITEMS_TABLE]' in para.text:
            table = doc.add_table(rows=1, cols=5)
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER # Center the whole table

            # Set column widths
            total_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            first_col_width = total_width * 0.5
            other_col_width = (total_width - first_col_width) / 4
            for j, col in enumerate(table.columns):
                if j == 0:
                    col.width = int(first_col_width)
                else:
                    col.width = int(other_col_width)

            # Style the table
            tbl = table._element
            tbl_pr = tbl.xpath('./w:tblPr')[0]

            # Full width
            table_width_element = OxmlElement('w:tblW')
            table_width_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', "0")
            table_width_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', "auto") # Use 'auto' for full width
            tbl_pr.append(table_width_element)

            # Remove left indent (if present).  This is the crucial part to fix the left spacing.
            try:
                table_ind = tbl_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblInd')
                if table_ind is not None:
                    tbl_pr.remove(table_ind)
            except Exception as e:
                print(f"Error removing table indent: {e}")

            # No gridlines
            table_borders_element = OxmlElement('w:tblBorders')
            table_borders_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top', "none")
            table_borders_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left', "none")
            table_borders_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom', "none")
            table_borders_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}right', "none")
            tbl_pr.append(table_borders_element)

            # Header row styling
            hdr_cells = table.rows[0].cells

            hdr_cells[0].text = 'Description'
            hdr_cells[1].text = 'Qty'
            hdr_cells[2].text = 'Unit Price'
            hdr_cells[3].text = params['table_vat_header']
            hdr_cells[4].text = f'Total ({data["currency"]})'
            for cn, cell in enumerate(hdr_cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    if cn == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    paragraph.style.paragraph_format.space_before = Pt(3)
                    paragraph.style.paragraph_format.space_after = Pt(3)
                    for run in paragraph.runs:
                        run.font.name = params['table_header_font']
                        run.font.size = Pt(params['table_header_fontsize'])
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd', {'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w': "clear", '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': "auto", '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill': params['table_header_fillcolor']}))

            pretax_total = Decimal(0.0)
            total = Decimal(0.0)

            for item in data['items']:
                item_subtotal = Decimal(0.0)
                item_total = Decimal(0.0)
                row_cells = table.add_row().cells
                row_cells[0].text = item['description']
                row_cells[1].text = str(item['quantity'])
                row_cells[2].text = str(item['unit_price'])
                row_cells[3].text = str(item['vat_pct']) + '%'

                item_subtotal = Decimal(item['quantity']) * Decimal(item['unit_price'])
                item_total = item_subtotal * Decimal((100.0 + item['vat_pct']) / 100.0)
                row_cells[4].text = invoicing.format_currency(item_total)
                pretax_total += item_subtotal
                total += item_total

                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for j, cell in enumerate(row_cells):
                    for paragraph in cell.paragraphs:
                        if j > 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        for run in paragraph.runs:
                            run.font.name = params['table_item_font']
                            run.font.size = Pt(params['table_item_fontsize'])
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = False

            # Add bottom border to the last row
            last_row = table.rows[-1]
            for cell in last_row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                bottom_border = OxmlElement('w:bottom', {'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'single', '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz': '4', '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': params['table_bottomline_color']})
                tc_borders.append(bottom_border)
                tc_pr.append(tc_borders)

            # Insert the table right after the placeholder paragraph
            p = para._element
            p.addnext(tbl)
            p.getparent().remove(p)
            break

    vat_total = total - pretax_total

    # Replace simple placeholders
    notes = data.get('notes', '')
    replacements = {
        '[INVOICE_NUMBER]': data['invoice_number'],
        '[DATE]': data['date'],
        '[DUE_DATE]': data['due_date'],
        '[CLIENT_NAME]': data['client_name'],
        '[CLIENT_ADDRESS]': data['client_address'],
        '[CLIENT_VAT]': data['client_vat_line'],
        '[CLIENT_EMAIL]': data.get('client_email', ''),
        '[SUBTOTAL]': invoicing.format_currency(pretax_total, data['currency'], cur1st),
        '[VAT_TOTAL]': invoicing.format_currency(vat_total, data['currency'], cur1st),
        '[TOTAL]': invoicing.format_currency(total, data['currency'], cur1st),
    }

    # handle credits, if any
    amount_due_total = total
    if data.get('credit'):
        amount_due_total -= Decimal(data['credit'])
        credit_fields = {
          '[CREDIT_TXT]': data['credit_text'],
          '[CREDIT]': invoicing.format_currency(-data['credit'], data['currency'], cur1st),
          '[AMOUNT_DUE_TXT]': data['amount_due_text'],
          '[FINAL_TOTAL]': invoicing.format_currency(amount_due_total, data['currency'], cur1st),
        }
        if amount_due_total <= 0:
            notes = data.get('no_payment_text', '') + notes
    else:
        credit_fields = {
          '[CREDIT_TXT]': '',
          '[CREDIT]': '',
          '[AMOUNT_DUE_TXT]': '',
          '[FINAL_TOTAL]': ''
        }
    replacements |= credit_fields
    replacements['[NOTES]'] = notes

    replace_placeholder(doc.paragraphs, replacements)

    doc.save(output_path)

    print(f"✅ Word-format invoice generated, with total: {RED}{total}{RESET}, amount due: {RED}{amount_due_total}{RESET}")


# TODO: May want to remove due date and instead include payment terms.

if __name__ == "__main__":
    billable_month = invoicing.get_relevant_month().strftime("%B")

    # Example usage
    invoice_data = {
        'invoice_number': 'INV-2025-003',
        'date': date.today().strftime('%Y-%m-%d'),
        'due_date': (date.today() + relativedelta(months=1)).strftime('%Y-%m-%d'),
        'client_name': 'Acme Corp',
        'client_address': '123 Main Street, Brussels, Belgium',
        'client_vat_line': 'VAT #: BE987654321',
        'client_email': 'info@acme.com',
        'currency': '€',
        'notes': 'This invoice is VAT exempt under Article 39 of the VAT Code. Copyright licensing and other terms are is pursuant to agreement dated 1 January 2020.',
        'items': [
            {'description': f'Consulting Services thru {billable_month} (hours)', 'quantity': 30, 'unit_price': 55, 'vat_pct': 0},
            {'description': 'Copyright Licensing (lump sum for delivery)', 'quantity': 1, 'unit_price': 1200, 'vat_pct': 21},
        ]
    }

    output_docx = f'invoice_{invoice_data["invoice_number"]}.docx'

    # Use US or EU template as applicable
    doc_params = {'table_header_fillcolor': '#FF00FF', 'currency_as_prefix': True}
    fill_word_invoice('templates/human_invoice_template.docx', output_docx, invoice_data, doc_params)

    print("Generated bill for", billable_month)
